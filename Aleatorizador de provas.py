import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import random
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt


def set_font(run):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Criando banco de questões para teste
numbers = list(range(1, 21))

subject = list()
item = list()
correct = list()
incorrect1 = list()
incorrect2 = list()
incorrect3 = list()
incorrect4 = list()

for op in ["adição", "subtração", "multiplicação", "divisão"]:
    for num1 in numbers:
        for num2 in numbers:
            if op == "adição":
                operator = "+"
                response = num1 + num2
            elif op == "subtração":
                operator = "–"
                response = num1 - num2
            elif op == "multiplicação":
                operator = "×"
                response = num1 * num2
            elif op == "divisão":
                operator = "/"
                response = num1 / num2
            
            correct.append(f"{response:.2f}")
            A = B = C = D = response
            subject.append(op)
            item.append(f"Quanto é {num1} {operator} {num2}?")
            
            for alternative, incorrects in zip([A, B, C, D], [incorrect1, incorrect2, incorrect3, incorrect4]):
                while True:
                    temp = random.choice(numbers)
                    if temp != correct and temp != A and temp!= B and temp != C and temp != D:
                        incorrects.append(f"{temp:.2f}")
                        break

# print("Done!")

# criando dataframe
df = pd.DataFrame({
     "tema": subject,
     'questão': item,
     'correta': correct,
     'incorreta_1': incorrect1,
     'incorreta_2': incorrect2,
     'incorreta_3': incorrect3,
     'incorreta_4': incorrect4}
)

# checando 5 primeiras e 5 últimas linhas do dataframe
# df.head(5)
# df.tail(5)

# salvando banco de dados para fins de teste
# df.to_csv("banco_de_questões.csv", index = False)

# Criando funções preliminares

# Função criar prova

# Criei função para aleatorizar provas, dado banco de dados (Df), temas e número de questões por tema
def criar_prova(df: pd.DataFrame, temas: list, num_questoes: list, shuffle = False):
    """
    Seleciona aleatoriamente um número específico de questões de cada tema para compor uma prova.
    
    Parâmetros:
    - df: DataFrame do Pandas com as colunas ['tema', 'questão', 'correta', 'incorreta1', 'incorreta2', 'incorreta3']
    - temas: Lista de temas a serem incluídos na prova
    - num_questoes: Lista do número de questões para cada tema
    - shuffle: Booleano que define se questões de diferentes temas serão embaralhadas (default: False)
    
    Retorna:
    - DataFrame com as questões selecionadas para a prova
    """

    provas = []

    for tema, num in zip(temas, num_questoes):
        # Filtra as questões do tema específico
        df_tema = df[df['tema'] == tema]
        
        # Verifica se há questões suficientes para o tema
        if len(df_tema) < num:
            num = len(df_tema)
        
        # Seleciona aleatoriamente 'num' questões do tema
        df_selecionadas = df_tema.sample(n=num)
        
        # Adiciona as questões selecionadas à lista de provas
        provas.append(df_selecionadas)
    
    # Concatena todas as questões selecionadas para formar a prova final
    prova_final = pd.concat(provas).reset_index(drop=True)
    
    if shuffle:
        # embaralha ordem das questões
        prova_final = prova_final.sample(frac = 1).reset_index(drop = True)
        
    prova_docx(prova_final)
    
    return prova_final


# Função prova_docx()

def prova_docx(prova_final):

    # cria documentos
    exam = Document()
    responses = Document()
   
    question = [("Instituição", "Qual é o nome da instituição de ensino?"),
                ("Curso", "Qual é o nome do curso?"),
                ("Disciplina", "Qual é o nome da disciplina?"),
                ("Professor(a)", "Qual é o seu nome?"),
                ("Prova", "Qual é o número da prova (insira um número inteiro)?")]
    
    exam_name = input("Digite o nome do arquivo a ser criado, sem o formato (ex: prova):\n")
        
    for i in range(len(question)):

        temp = input(f"{question[i][1]} ")

        for file in [exam, responses]:

            if i == (len(question) - 1):

                file.add_paragraph("")
                if file == exam:
                    heading = file.add_heading(f"{question[i][0]}: {temp}", level = 1)
                else:
                    heading = file.add_heading(f"{question[i][0]}: {temp} – Gabarito", level = 1)
                run = heading.runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centralizar título de nível 1
                set_font(run)
                file.add_paragraph("")

            else:
                head = file.add_paragraph()
                run = head.add_run(f"{question[i][0]}: {temp}")
                run.bold = True
                

    exam.add_paragraph("Nome completo: ___________________________________________________________________________________")
    exam.add_paragraph("Matrícula: ____________________________________________ Data: __________/__________/_______________")
    exam.add_paragraph("")

    # preenchendo prova e gabarito
    for i, row in enumerate(prova_final.itertuples(index = True, name = "Pandas"), 1):
        # PROVA
        
        # armazena alternativas
        auxiliar = [row.correta, row.incorreta_1, row.incorreta_2, row.incorreta_3, row.incorreta_4]
        # e depois as embaralha
        random.shuffle(auxiliar)

        # enunciado da questão
        p = exam.add_paragraph()
        run = p.add_run(f"Questão {i}. {row.questão}")  # Escreve o enunciado da questão em negrito
        run.bold = True

        # Adiciona as alternativas
        exam.add_paragraph(f"(     ) A. {auxiliar[0]}.")
        exam.add_paragraph(f"(     ) B. {auxiliar[1]}.")
        exam.add_paragraph(f"(     ) C. {auxiliar[2]}.")
        exam.add_paragraph(f"(     ) D. {auxiliar[3]}.")
        exam.add_paragraph(f"(     ) E. {auxiliar[4]}.")
        exam.add_paragraph("")
        
        # GABARITO
        if i == 1:
            table = responses.add_table(rows = 1, cols = 3)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Questão'
            hdr_cells[1].text = 'Gabarito'
            hdr_cells[2].text = 'Tema'
            
        # Negrita e centraliza o texto nas células do cabeçalho
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        
        # checando posição do gabarito nas alternativas
        positions = ["A", "B", "C", "D", "E"]
        if row.correta in auxiliar:
            pos = auxiliar.index(row.correta)
        
        row_cells[1].text = f"{positions[pos]}"
        row_cells[2].text = row.tema
        
        # Centraliza o texto nas células
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
        
    exam.save(f"{exam_name}.docx") 
    responses.save(f"{exam_name}_gabarito.docx")

# print("Done!")

# Teste do programa
# exemplo: 10 questões, sendo 2 do tema A, 3 do tema B, 4 do tema C e 1 do tema D
questions = criar_prova(df = df,
                        temas = ["adição", "subtração", "multiplicação", "divisão"],
                        num_questoes = [5, 4, 3, 3],
                        shuffle = True) 

questions.head(10)




